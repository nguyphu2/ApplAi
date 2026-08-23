import io
import json
from unittest.mock import MagicMock

from docx import Document


def make_event(user_id='user-123', body=None):
    event = {
        'requestContext': {
            'authorizer': {'jwt': {'claims': {'sub': user_id}}},
        },
    }
    if body is not None:
        event['body'] = json.dumps(body)
    return event


def make_test_docx(paragraphs):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def read_docx_paragraphs(docx_bytes):
    document = Document(io.BytesIO(docx_bytes))
    return [p.text for p in document.paragraphs]


def make_tool_response(tool_name, input_data):
    return {
        'body': MagicMock(read=lambda: json.dumps({
            'content': [{'type': 'tool_use', 'name': tool_name, 'input': input_data}],
        }).encode('utf-8'))
    }


def make_three_stage_responses(
    missing_keywords=('Flask', 'Kubernetes', 'CI/CD', 'GraphQL', 'Terraform'),
    red_flags=('No quantified impact', 'Vague job titles', 'Missing recent tech stack'),
    match_score_before=40,
    rewrite_rewrites=(),
    scan_rewrites=(),
    match_score_after=80,
):
    return [
        make_tool_response('report_analysis', {
            'match_score_before': match_score_before,
            'missing_keywords': list(missing_keywords),
            'red_flags': list(red_flags),
        }),
        make_tool_response('report_rewrite', {'rewrites': list(rewrite_rewrites)}),
        make_tool_response('report_scan', {'rewrites': list(scan_rewrites), 'match_score_after': match_score_after}),
    ]


def upload_test_resume(handler_module, user_id, resume_id, paragraphs, filename='resume.docx'):
    docx_bytes = make_test_docx(paragraphs)
    handler_module.s3_client.put_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key(user_id, resume_id, 'docx'),
        Body=docx_bytes,
    )
    handler_module.table.put_item(Item={
        'user_id': user_id,
        'resumes': [{
            'id': resume_id,
            'filename': filename,
            'text': '\n'.join(paragraphs),
            'file_type': 'docx',
            'uploaded_at': '2026-01-01T00:00:00Z',
        }],
        'active_resume_ids': [resume_id],
    })


def test_rejects_when_job_description_missing(handler_module):
    response = handler_module.handler(
        make_event(body={'resume_id': 'r1', 'job_description_text': ''}), None
    )

    assert response['statusCode'] == 400


def test_rejects_when_resume_not_found(handler_module):
    handler_module.table.put_item(Item={'user_id': 'user-123', 'resumes': []})

    response = handler_module.handler(
        make_event(body={'resume_id': 'missing', 'job_description_text': 'Looking for a Python engineer'}), None
    )

    assert response['statusCode'] == 400


def test_rejects_when_resume_is_not_docx(handler_module):
    handler_module.table.put_item(Item={
        'user_id': 'user-123',
        'resumes': [{'id': 'r1', 'filename': 'resume.pdf', 'text': 'old text', 'uploaded_at': '2026-01-01T00:00:00Z'}],
    })

    response = handler_module.handler(
        make_event(body={'resume_id': 'r1', 'job_description_text': 'Looking for a Python engineer'}), None
    )

    assert response['statusCode'] == 400


def test_overwrite_runs_three_stage_chain_and_saves_result(handler_module, monkeypatch):
    upload_test_resume(handler_module, 'user-123', 'r1', ['Ada Lovelace', 'Software Engineer', 'Built REST APIs in Python'])
    invoke_model = MagicMock(side_effect=make_three_stage_responses(
        rewrite_rewrites=[{'paragraph_index': 2, 'new_text': 'Built REST APIs in Python and Flask, cutting response time by 30%'}],
        match_score_after=85,
    ))
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for a Python and Flask engineer',
        'target_match_percent': 80,
        'one_page': True,
        'save_as_new_copy': False,
    }), None)

    assert response['statusCode'] == 200
    body = json.loads(response['body'])
    assert body['match_score_before'] == 40
    assert body['match_score_after'] == 85
    assert body['missing_keywords'] == ['Flask', 'Kubernetes', 'CI/CD', 'GraphQL', 'Terraform']
    assert body['red_flags'] == ['No quantified impact', 'Vague job titles', 'Missing recent tech stack']
    assert body['resume_id'] == 'r1'
    assert body['filename'] == 'resume.docx'
    assert invoke_model.call_count == 3

    stored = handler_module.s3_client.get_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key('user-123', 'r1', 'docx'),
    )
    paragraphs = read_docx_paragraphs(stored['Body'].read())
    assert paragraphs[0] == 'Ada Lovelace'
    assert paragraphs[2] == 'Built REST APIs in Python and Flask, cutting response time by 30%'

    item = handler_module.table.get_item(Key={'user_id': 'user-123'})['Item']
    assert len(item['resumes']) == 1
    assert 'cutting response time by 30%' in item['resumes'][0]['text']


def test_scan_stage_rewrites_apply_on_top_of_rewrite_stage(handler_module, monkeypatch):
    upload_test_resume(handler_module, 'user-123', 'r1', ['Ada Lovelace', 'Built REST APIs', 'Led a team of 3'])
    invoke_model = MagicMock(side_effect=make_three_stage_responses(
        rewrite_rewrites=[{'paragraph_index': 1, 'new_text': 'Built REST APIs in Python and Flask'}],
        scan_rewrites=[{'paragraph_index': 2, 'new_text': 'Led a team of 3 engineers, shipping 4 features per quarter'}],
    ))
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for an engineering lead',
    }), None)

    assert response['statusCode'] == 200
    stored = handler_module.s3_client.get_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key('user-123', 'r1', 'docx'),
    )
    paragraphs = read_docx_paragraphs(stored['Body'].read())
    assert paragraphs[1] == 'Built REST APIs in Python and Flask'
    assert paragraphs[2] == 'Led a team of 3 engineers, shipping 4 features per quarter'


def test_save_as_new_copy_creates_new_entry_without_touching_original(handler_module, monkeypatch):
    upload_test_resume(handler_module, 'user-123', 'r1', ['Ada Lovelace', 'Software Engineer', 'Built REST APIs in Python'])
    invoke_model = MagicMock(side_effect=make_three_stage_responses(
        rewrite_rewrites=[{'paragraph_index': 2, 'new_text': 'Built scalable REST APIs in Python and Flask'}],
    ))
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for a Python and Flask engineer',
        'save_as_new_copy': True,
    }), None)

    assert response['statusCode'] == 200
    body = json.loads(response['body'])
    assert body['resume_id'] != 'r1'
    assert body['filename'] == 'resume (optimized).docx'

    item = handler_module.table.get_item(Key={'user_id': 'user-123'})['Item']
    assert len(item['resumes']) == 2
    original = next(r for r in item['resumes'] if r['id'] == 'r1')
    assert original['text'] == 'Ada Lovelace\nSoftware Engineer\nBuilt REST APIs in Python'

    original_stored = handler_module.s3_client.get_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key('user-123', 'r1', 'docx'),
    )
    assert read_docx_paragraphs(original_stored['Body'].read())[2] == 'Built REST APIs in Python'


def test_bedrock_failure_returns_502(handler_module, monkeypatch):
    upload_test_resume(handler_module, 'user-123', 'r1', ['Ada Lovelace', 'Software Engineer', 'Built REST APIs'])
    invoke_model = MagicMock(side_effect=Exception('bedrock unavailable'))
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for an engineer',
    }), None)

    assert response['statusCode'] == 502


def test_second_stage_failure_leaves_original_resume_untouched(handler_module, monkeypatch):
    upload_test_resume(handler_module, 'user-123', 'r1', ['Ada Lovelace', 'Software Engineer', 'Built REST APIs'])
    invoke_model = MagicMock(side_effect=[
        make_tool_response('report_analysis', {
            'match_score_before': 40,
            'missing_keywords': ['Flask', 'Kubernetes', 'CI/CD', 'GraphQL', 'Terraform'],
            'red_flags': ['No quantified impact', 'Vague job titles', 'Missing recent tech stack'],
        }),
        Exception('bedrock unavailable'),
    ])
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for an engineer',
    }), None)

    assert response['statusCode'] == 502
    stored = handler_module.s3_client.get_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key('user-123', 'r1', 'docx'),
    )
    assert read_docx_paragraphs(stored['Body'].read()) == ['Ada Lovelace', 'Software Engineer', 'Built REST APIs']
    item = handler_module.table.get_item(Key={'user_id': 'user-123'})['Item']
    assert item['resumes'][0]['text'] == 'Ada Lovelace\nSoftware Engineer\nBuilt REST APIs'


def test_near_empty_resume_returns_400(handler_module, monkeypatch):
    # A table-based resume (or any near-empty extraction) should fail fast
    # with a clear error rather than proceeding to burn three Bedrock calls
    # on an essentially blank document.
    upload_test_resume(handler_module, 'user-123', 'r1', ['Ada Lovelace'])
    invoke_model = MagicMock(side_effect=make_three_stage_responses())
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for an engineer',
    }), None)

    assert response['statusCode'] == 400
    body = json.loads(response['body'])
    assert 'table-based layout' in body['error']
    assert invoke_model.call_count == 0


def make_multi_run_test_docx():
    document = Document()
    document.add_paragraph('Ada Lovelace')
    document.add_paragraph('Software Engineer')
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run('Built REST APIs')
    run1.bold = True
    paragraph.add_run(' using Python')
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def upload_multi_run_test_resume(handler_module, user_id, resume_id, filename='resume.docx'):
    docx_bytes = make_multi_run_test_docx()
    handler_module.s3_client.put_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key(user_id, resume_id, 'docx'),
        Body=docx_bytes,
    )
    handler_module.table.put_item(Item={
        'user_id': user_id,
        'resumes': [{
            'id': resume_id,
            'filename': filename,
            'text': 'Built REST APIs using Python',
            'file_type': 'docx',
            'uploaded_at': '2026-01-01T00:00:00Z',
        }],
        'active_resume_ids': [resume_id],
    })


def test_multi_run_paragraph_rewrite_preserves_formatting_and_clears_stray_runs(handler_module, monkeypatch):
    upload_multi_run_test_resume(handler_module, 'user-123', 'r1')
    invoke_model = MagicMock(side_effect=make_three_stage_responses(
        rewrite_rewrites=[{'paragraph_index': 2, 'new_text': 'Shipped scalable backend services'}],
    ))
    monkeypatch.setattr(handler_module, 'bedrock_runtime', MagicMock(invoke_model=invoke_model))

    response = handler_module.handler(make_event(body={
        'resume_id': 'r1',
        'job_description_text': 'Looking for a backend engineer',
    }), None)

    assert response['statusCode'] == 200
    stored = handler_module.s3_client.get_object(
        Bucket=handler_module.BUCKET_NAME,
        Key=handler_module.resume_file_key('user-123', 'r1', 'docx'),
    )
    document = Document(io.BytesIO(stored['Body'].read()))
    paragraph = document.paragraphs[2]

    assert paragraph.text == 'Shipped scalable backend services'
    assert paragraph.runs[0].bold is True
    assert 'using Python' not in paragraph.text

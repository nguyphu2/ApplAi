"""Triggered by API Gateway, behind a Cognito JWT authorizer. Rewrites a
saved DOCX resume's bullet points/summary toward a target keyword-match
percentage against a job description, editing the original document's
paragraph runs in place so every font, margin, and layout choice from the
original file survives untouched - regenerating a document from scratch
never reproduces its original visual design.

Runs three sequential Bedrock tool-use calls per request rather than one
call doing everything: an "analyze" pass (recruiter persona) scores the
resume and finds gaps, a "rewrite" pass (XYZ formula) closes them, and a
"scan" pass (ATS/hiring-manager persona) polishes what's left and scores
the final result. If any of the three fails, nothing is written to S3 or
DynamoDB - the resume stays exactly as it was.
"""
import io
import json
import os
import uuid
from datetime import datetime, timezone

import boto3
from botocore.exceptions import ClientError
from docx import Document
from dotenv import load_dotenv

load_dotenv('.env')
TABLE_NAME = os.getenv('SETTINGS_TABLE_NAME')
BUCKET_NAME = os.getenv('BUCKET_NAME')
CLAUDE_MODEL_ID = 'us.anthropic.claude-haiku-4-5-20251001-v1:0'

dynamodb = boto3.resource('dynamodb')
table = dynamodb.Table(TABLE_NAME)
s3_client = boto3.client('s3')
bedrock_runtime = boto3.client('bedrock-runtime')

REWRITE_ITEM_SCHEMA = {
    'type': 'object',
    'properties': {
        'paragraph_index': {'type': 'integer'},
        'new_text': {'type': 'string'},
    },
    'required': ['paragraph_index', 'new_text'],
}

ANALYZE_TOOL_SCHEMA = {
    'type': 'object',
    'properties': {
        'match_score_before': {'type': 'integer'},
        'missing_keywords': {'type': 'array', 'items': {'type': 'string'}},
        'red_flags': {'type': 'array', 'items': {'type': 'string'}},
    },
    'required': ['match_score_before', 'missing_keywords', 'red_flags'],
}

REWRITE_TOOL_SCHEMA = {
    'type': 'object',
    'properties': {'rewrites': {'type': 'array', 'items': REWRITE_ITEM_SCHEMA}},
    'required': ['rewrites'],
}

SCAN_TOOL_SCHEMA = {
    'type': 'object',
    'properties': {
        'rewrites': {'type': 'array', 'items': REWRITE_ITEM_SCHEMA},
        'match_score_after': {'type': 'integer'},
    },
    'required': ['rewrites', 'match_score_after'],
}


def get_user_id(event):
    return event['requestContext']['authorizer']['jwt']['claims']['sub']


def resume_file_key(user_id, resume_id, file_type='docx'):
    extension = 'docx' if file_type == 'docx' else 'pdf'
    return f'resumes/{user_id}/{resume_id}.{extension}'


def extract_paragraphs(docx_bytes):
    document = Document(io.BytesIO(docx_bytes))
    return document, [p.text for p in document.paragraphs]


def indexed_paragraphs_json(paragraphs):
    indexed = [{'paragraph_index': i, 'text': t} for i, t in enumerate(paragraphs) if t.strip()]
    return json.dumps(indexed)


def build_analyze_prompt(paragraphs, job_description_text):
    return f"""Act as a senior recruiter for this exact company. Analyze this
resume against this job description.

Job posting:
{job_description_text}

Resume paragraphs:
{indexed_paragraphs_json(paragraphs)}

Give a match score out of 100 (match_score_before), the top 5 keywords
from the job posting that are missing from the resume (missing_keywords),
and the 3 red flags a hiring manager would spot in this resume in under
10 seconds (red_flags) - things like unquantified claims, vague titles,
or an obviously missing must-have skill."""


def build_rewrite_prompt(paragraphs, job_description_text, missing_keywords, red_flags, one_page):
    length_instruction = (
        'Keep each rewritten paragraph roughly the same length as the '
        'original - this resume must still fit on one page, so '
        'significantly longer replacement text risks pushing content onto '
        'a second page.'
        if one_page else
        'Rewritten paragraphs do not need to match the original length.'
    )
    return f"""Rewrite this resume's experience section to naturally
include these keywords and remove these red flags, using the formula
"Accomplished X, as measured by Y, by doing Z" wherever the underlying
achievement supports it.

Job posting:
{job_description_text}

Missing keywords to naturally include: {json.dumps(missing_keywords)}
Red flags to remove: {json.dumps(red_flags)}

Resume paragraphs (each has a paragraph_index and its current text):
{indexed_paragraphs_json(paragraphs)}

Only rephrase and re-emphasize content that is already genuinely present
in the resume - never invent experience, skills, credentials, metrics, or
achievements that aren't already there. {length_instruction} Only include
a paragraph in "rewrites" if you are actually changing its text - omit
any paragraph you are leaving as-is."""


def build_scan_prompt(paragraphs, job_description_text, target_match_percent, one_page):
    length_instruction = (
        'Keep each rewritten paragraph roughly the same length as the '
        'original - this resume must still fit on one page.'
        if one_page else
        'Rewritten paragraphs do not need to match the original length.'
    )
    return f"""Now act as an ATS filter and a hiring manager reading 200
resumes in one sitting. Scan this resume (already rewritten once) and
identify which sections would still get skipped over, then rewrite just
those so they stop the scroll.

Job posting:
{job_description_text}

Resume paragraphs (each has a paragraph_index and its current text):
{indexed_paragraphs_json(paragraphs)}

Only rephrase and re-emphasize content that is already genuinely present
in the resume - never invent experience, skills, credentials, metrics, or
achievements that aren't already there. {length_instruction} Only include
a paragraph in "rewrites" if you are actually changing its text - omit
any paragraph you are leaving as-is.

Also assess this resume's resulting keyword match against the job posting
as match_score_after (0-100). Aim for around {target_match_percent}%, but
honesty always takes priority over hitting that number: if it isn't
reachable without fabricating, get as close as truthfully possible and
report the real score, which may fall short of the target."""


def invoke_claude_tool(prompt, tool_name, tool_schema, max_tokens):
    response = bedrock_runtime.invoke_model(
        modelId=CLAUDE_MODEL_ID,
        body=json.dumps({
            'anthropic_version': 'bedrock-2023-05-31',
            'max_tokens': max_tokens,
            'tools': [{
                'name': tool_name,
                'description': f'Report the {tool_name.replace("_", " ")} result.',
                'input_schema': tool_schema,
            }],
            'tool_choice': {'type': 'tool', 'name': tool_name},
            'messages': [{'role': 'user', 'content': prompt}],
        }),
    )
    result = json.loads(response['body'].read())
    tool_use = next(block for block in result['content'] if block.get('type') == 'tool_use')
    return tool_use['input']


def apply_rewrites(document, rewrites):
    for rewrite in rewrites:
        index = rewrite.get('paragraph_index')
        new_text = rewrite.get('new_text')
        if not isinstance(index, int) or not isinstance(new_text, str):
            continue
        if index < 0 or index >= len(document.paragraphs):
            continue
        paragraph = document.paragraphs[index]
        if not paragraph.runs:
            continue
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def handler(event, context):
    user_id = get_user_id(event)
    body = json.loads(event.get('body') or '{}')

    resume_id = body.get('resume_id')
    job_description_text = (body.get('job_description_text') or '').strip()
    job_description_text = job_description_text[:6000]
    target_match_percent = body.get('target_match_percent', 70)
    one_page = bool(body.get('one_page'))
    save_as_new_copy = bool(body.get('save_as_new_copy'))

    if not job_description_text:
        return {'statusCode': 400, 'body': json.dumps({'error': 'job description text is required'})}

    try:
        response = table.get_item(Key={'user_id': user_id})
    except ClientError as e:
        print(f'DynamoDB get_item failed: {e}')
        return {'statusCode': 502, 'body': json.dumps({'error': 'optimize service failed'})}

    item = response.get('Item', {})
    resumes = item.get('resumes', [])
    resume = next((r for r in resumes if r['id'] == resume_id), None)
    if resume is None or resume.get('file_type') != 'docx':
        return {'statusCode': 400, 'body': json.dumps({'error': 'resume must be a DOCX to optimize'})}

    try:
        s3_object = s3_client.get_object(Bucket=BUCKET_NAME, Key=resume_file_key(user_id, resume_id, 'docx'))
        docx_bytes = s3_object['Body'].read()
    except ClientError as e:
        print(f'could not download resume DOCX: {e}')
        return {'statusCode': 502, 'body': json.dumps({'error': 'optimize service failed'})}

    try:
        document, paragraphs = extract_paragraphs(docx_bytes)
    except Exception as e:
        print(f'could not parse resume DOCX: {e}')
        return {'statusCode': 502, 'body': json.dumps({'error': 'could not read resume DOCX'})}

    non_empty_count = sum(1 for p in paragraphs if p.strip())
    if non_empty_count < 3:
        return {'statusCode': 400, 'body': json.dumps({'error': 'resume has too little extractable text - it may use a table-based layout that is not yet supported'})}

    try:
        analysis = invoke_claude_tool(
            build_analyze_prompt(paragraphs, job_description_text),
            'report_analysis', ANALYZE_TOOL_SCHEMA, max_tokens=1000,
        )
        match_score_before = analysis.get('match_score_before', 0)
        missing_keywords = analysis.get('missing_keywords', [])
        red_flags = analysis.get('red_flags', [])

        rewrite_result = invoke_claude_tool(
            build_rewrite_prompt(paragraphs, job_description_text, missing_keywords, red_flags, one_page),
            'report_rewrite', REWRITE_TOOL_SCHEMA, max_tokens=2000,
        )
        apply_rewrites(document, rewrite_result.get('rewrites', []))
        paragraphs = [p.text for p in document.paragraphs]

        scan_result = invoke_claude_tool(
            build_scan_prompt(paragraphs, job_description_text, target_match_percent, one_page),
            'report_scan', SCAN_TOOL_SCHEMA, max_tokens=2000,
        )
        apply_rewrites(document, scan_result.get('rewrites', []))
        match_score_after = scan_result.get('match_score_after', match_score_before)
    except Exception as e:
        print(f'optimize failed: {e}')
        return {'statusCode': 502, 'body': json.dumps({'error': 'optimize failed'})}

    output_buffer = io.BytesIO()
    document.save(output_buffer)
    new_docx_bytes = output_buffer.getvalue()
    new_text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())

    if save_as_new_copy:
        new_resume_id = str(uuid.uuid4())
        original_filename = resume.get('filename', 'resume.docx')
        if '.' in original_filename:
            base, ext = original_filename.rsplit('.', 1)
            new_filename = f'{base} (optimized).{ext}'
        else:
            new_filename = f'{original_filename} (optimized)'
        try:
            s3_client.put_object(
                Bucket=BUCKET_NAME,
                Key=resume_file_key(user_id, new_resume_id, 'docx'),
                Body=new_docx_bytes,
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        except ClientError as e:
            print(f'could not upload optimized resume: {e}')
            return {'statusCode': 502, 'body': json.dumps({'error': 'optimize service failed'})}
        resumes.append({
            'id': new_resume_id,
            'filename': new_filename,
            'text': new_text,
            'file_type': 'docx',
            'uploaded_at': datetime.now(timezone.utc).isoformat(),
        })
        result_resume_id = new_resume_id
        result_filename = new_filename
    else:
        try:
            s3_client.put_object(
                Bucket=BUCKET_NAME,
                Key=resume_file_key(user_id, resume_id, 'docx'),
                Body=new_docx_bytes,
                ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        except ClientError as e:
            print(f'could not upload optimized resume: {e}')
            return {'statusCode': 502, 'body': json.dumps({'error': 'optimize service failed'})}
        resume['text'] = new_text
        result_resume_id = resume_id
        result_filename = resume.get('filename', 'resume.docx')

    try:
        table.put_item(Item={**item, 'user_id': user_id, 'resumes': resumes})
    except ClientError as e:
        print(f'DynamoDB put_item failed: {e}')
        return {'statusCode': 502, 'body': json.dumps({'error': 'optimize service failed'})}

    return {
        'statusCode': 200,
        'body': json.dumps({
            'match_score_before': match_score_before,
            'match_score_after': match_score_after,
            'missing_keywords': missing_keywords,
            'red_flags': red_flags,
            'resume_id': result_resume_id,
            'filename': result_filename,
        }),
    }
